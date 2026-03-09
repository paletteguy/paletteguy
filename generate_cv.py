#!/usr/bin/env python3

import os
import io
import subprocess
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ============================================================
# CV DATA — single source of truth for .docx, .pdf, and .md
# ============================================================

CV_DATA = {
    'name': 'Karsten Sperling Opdal',
    'tagline': 'Developing Embedded Systems into the future of AI',
    'photo': 'me.png',
    'contact': [
        ('Company:', 'Opdal Enterprise'),
        ('CVR:', '31357667'),
        ('Address:', 'Nordvangen 13, DK-4600 Køge'),
        ('Phone:', '+45 42 40 42 82'),
        ('Email:', 'karsten.s.opdal@gmail.com'),
        ('LinkedIn:', 'https://linkedin.com/in/karstenopdal-7a11512'),
        ('Web:', 'https://opdal.dk'),
        ('GitHub profile:', 'https://github.com/paletteguy'),
    ],
    'summary': (
        "I make embedded Linux do what it's told \u2014 on hardware where failure isn't an option.\n\n"
        "For 15 years at Laerdal Medical, I was the architect behind SimPad's embedded platform \u2014 owning "
        "everything from the Yocto build system delivering production Linux across ARM and x86, to the real-time "
        "simulation engine that brings medical manikins to life via CAN bus, IOC controllers, and physiological "
        "modelling. I evolved the build system through six major Yocto releases (Rocko to Scarthgap), migrated "
        "the platform from Ubuntu to a custom distribution with RAUC A/B OTA updates and HSM-signed security, "
        "and architected CI/CD pipelines with SBOM generation and CVE tracking for medical device compliance. "
        "Along the way I built Qt/QML touch-screen UIs, ALSA audio pipelines, WiFi/Bluetooth integration "
        "(with patches contributed to BlueZ and Qt Bluetooth), and a suite of internal tools \u2014 from an "
        "AI-powered diagnostic platform to an SSH certificate portal and a cross-platform device imager.\n\n"
        "Before Laerdal, I spent a decade in the Nokia ecosystem \u2014 developing the user interfaces for the "
        "iconic 3210 and 3310, implementing concatenated SMS and T9 input, and building simulation tools that "
        "let engineers develop phone software in Visual Studio without touching hardware. I debugged firmware "
        "issues nobody else could crack, including a battery drain bug that improved battery lifetime by 30%. "
        "At Infineon, I shipped the Nokia 2110 low-cost platform where every kilobyte counted.\n\n"
        "I own the full stack from silicon to screen: board bring-up on i.MX6, i.MX8, and x86, Linux kernel "
        "debugging and Device Tree customization, C++ services in real-time, and secure OTA deployment. "
        "I've integrated AI deeply into my engineering workflow \u2014 from AI-assisted code generation and review "
        "to building custom tooling around LLMs for debugging and documentation.\n\n"
        "30+ years from Nokia handsets to medical devices. Now looking for the next platform worth building.\n\n"
        "Copenhagen, Denmark | Open to remote across Europe"
    ),
    'skills': (
        'C++ (C++23) \u00b7 Embedded Linux \u00b7 Yocto \u00b7 Linux kernel \u00b7 Qt/QML (Qt 3\u20136) \u00b7 Python \u00b7 '
        'WiFi/Bluetooth \u00b7 CI/CD \u00b7 Boost \u00b7 AI-assisted engineering'
    ),
    'languages': [
        ('Danish', 'Native speaker (written and spoken)'),
        ('English', 'Fluent \u2014 second language used throughout career, full professional proficiency'),
        ('Norwegian', "Fluent understanding (spoken and written), conversational speaking \u2014 25 years' exposure, Norwegian spouse"),
        ('Swedish', "Strong understanding (spoken and written), near-conversational speaking \u2014 30 years' exposure, neighbouring country"),
        ('German', 'Good understanding (spoken and written), speaking somewhat rusty \u2014 10 years during education'),
    ],
    'experience': [
        {
            'company': 'Laerdal Medical',
            'title': 'Senior Software Engineer Consulting',
            'period': 'October 2010 \u2013 February 2026 (15 years 5 months)',
            'bullets': [
                'Senior Embedded Software Engineer / Build System Architect for SimPad devices / simulators and SimMan3G product lines.',
                'Lead maintainer of Yocto/OpenEmbedded build platform delivering production Linux across ARM32, ARM64, x86-32, and x86-64. Evolved through six major Yocto releases (Rocko to Scarthgap). Migrated from Ubuntu to custom distribution with RAUC A/B OTA updates and Azure Key Vault HSM signing.',
                'Architected multi-platform CI/CD pipelines with GitHub Actions, SBOM generation, and automated CVE tracking for medical device compliance.',
                'Linux kernel configuration, debugging, patching, and Device Tree customization across i.MX6, i.MX8M Plus, and Intel x86 \u2014 board bring-up, peripheral enablement, pin muxing, clock trees, and driver troubleshooting.',
                'Developed real-time simulation engine controlling medical manikins (SimMom, SimBaby, MammaAnne, SimMan ALS) via CAN bus or IOC controller and physiological modelling. C++ (up to C++23) with Boost, Qt 3\u20136 with deep QML expertise for embedded UIs. Proficient in C#.',
                'ALSA audio pipelines for clinical simulation. Bluetooth expert with patches to BlueZ and Qt Bluetooth stacks.',
                'Created SimServer Imager (Qt6/C++/QML) for managed WIC/VSI/SPU deployment and CDN distribution. Built VEX Kernel Checker \u2014 AI-assisted CVE analysis integrated with Dependency-Track.',
                'Built SSH Certificate Portal (FastAPI/Python) \u2014 self-service time-limited SSH certificates with Azure AD OIDC, HSM-protected CA signing (RSA/ED25519/ECDSA), deployed on Azure App Service.',
                'Designed Azure Key Vault HSM signing workflow for RAUC bundles and secure boot \u2014 non-exportable keys, automated renewal via Azure Automation, multi-environment CI integration.',
                'Built Jira Analyse Companion \u2014 AI-powered diagnostic platform using Claude/Gemini/OpenAI for crash analysis across SimPad, LinkBox, and CAN firmware. Jira automation, source context integration. Delivered as CLI, VS Code extension, and Tauri desktop app.',
            ],
            'page_break_after': True,
        },
        {
            'company': 'Oscilloscope',
            'title': 'Senior Software Engineer Consulting',
            'period': 'October 2010 \u2013 June 2019 (8 years 9 months)',
            'bullets': ['Freelance'],
        },
        {
            'company': 'Nokia',
            'title': 'Software Engineering Specialist',
            'period': 'April 2010 \u2013 October 2010 (7 months)',
            'bullets': [
                'I excelled in identifying and resolving complex firmware issues at Nokia, significantly enhancing product performance.',
                'Specialized in debugging intricate hardware/software interaction bugs on mobile platforms.',
                'Successfully fixed a battery drain issue that had persisted for years, improving battery lifetime by 30%.',
                'Developed critical skills in embedded systems analysis and problem-solving within a leading technology company.',
            ],
        },
        {
            'company': 'Nokia Mobile Phones',
            'title': 'Senior Software Engineering Consultant',
            'period': 'September 2008 \u2013 February 2010 (1 year 6 months)',
            'bullets': [
                'Provided expert consultation on the S40 platform software, focusing on bug fixing and error correction.',
                'Debugged firmware issues across the S40 software stack to enhance performance and reliability.',
                'Collaborated with cross-functional teams to ensure timely resolution of software defects, improving user experience.',
            ],
            'page_break_after': True,
        },
        {
            'company': 'Infineon',
            'title': 'Software Expert',
            'period': 'October 2007 \u2013 September 2008 (1 year)',
            'bullets': [
                'Developed the Nokia 2110 low-cost phone platform during an expat assignment in Copenhagen.',
                'Optimized embedded software for resource-constrained hardware, ensuring efficient use of memory.',
                'Delivered high-quality software under tight deadlines, focusing on minimizing resource usage.',
            ],
        },
        {
            'company': 'Thors\u00f8 Data',
            'title': 'Software Expert',
            'period': 'October 2007 \u2013 September 2008 (1 year)',
            'bullets': ['In house Senior Contract Software Engineer'],
        },
        {
            'company': 'Tang-Data A/S',
            'title': 'Senior Software Engineer',
            'period': 'March 2007 \u2013 October 2007 (8 months)',
            'bullets': [
                'Developed a comprehensive veterinary CRM system utilizing Qt 3 and Qt 4, enhancing client management.',
                'Provided on-site support for veterinary computer setups, ensuring seamless CRM hardware installations.',
                'Collaborated with cross-functional teams to address customer needs and improve system functionality.',
            ],
            'page_break_after': True,
        },
        {
            'company': 'Nokia',
            'title': 'Senior Software Engineer',
            'period': 'October 2000 \u2013 February 2007 (6 years 5 months)',
            'bullets': [
                'Enhanced the software development lifecycle at Nokia through innovative UI development.',
                'Developed internal tools that streamlined handset software testing processes.',
                'Enabled mobile phone development in a simulated environment using Microsoft Visual Studio.',
            ],
        },
        {
            'company': 'EC-Soft Danmark A/S',
            'title': 'Senior System Software Engineer',
            'period': 'October 2000 \u2013 December 2001 (1 year 3 months)',
            'bullets': ['Consultant work for Nokia Denmark A/S developing S30 phones'],
        },
        {
            'company': 'Telenor',
            'title': 'Contract Software Engineer',
            'period': 'April 2000 \u2013 June 2000 (3 months)',
            'bullets': ['Developed a web service for read outlook mails and calendar on phones'],
        },
        {
            'company': 'EC-Soft Norge AS',
            'title': 'Senior Software Engineer',
            'period': 'May 2000 \u2013 October 2000 (6 months)',
            'bullets': ['In house Software Consultant'],
            'page_break_after': True,
        },
        {
            'company': 'Nokia',
            'title': 'Contract Software Engineer',
            'period': 'May 1999 \u2013 March 2000 (11 months)',
            'bullets': [
                "Developed user interfaces for Nokia's iconic handset models, including the 3210 and 3310.",
                'Collaborated with a fellow developer to implement concatenated SMS messaging, enhancing user communication.',
                'Integrated T9 input support, improving text input efficiency for users.',
            ],
        },
        {
            'company': 'EC-Soft Danmark A/S',
            'title': 'Software Consultant',
            'period': 'May 1999 \u2013 March 2000 (11 months)',
            'bullets': ['In house software consultant'],
        },
        {
            'company': 'Vizion Factory ApS',
            'title': 'Software Engineer',
            'period': 'October 1995 \u2013 April 1999 (3 years 7 months)',
            'bullets': [
                'Played a key role in software development and early web solutions, contributing to the burgeoning internet landscape.',
                'Built applications that addressed emerging needs and developed a training engine paired with a web interface to streamline user interaction.',
            ],
        },
        {
            'company': 'Greve Kommune',
            'title': 'Network System Specialist',
            'period': 'January 1991 \u2013 October 1995 (4 years 10 months)',
            'bullets': [
                'Set up new networks to enhance connectivity and efficiency within the organization.',
                'Provided training to employees on computer use, fostering a tech-savvy workplace.',
                'Maintained hardware by adding and replacing components, ensuring optimal performance.',
            ],
            'page_break_after': True,
        },
        {
            'company': 'Opdal Enterprise',
            'title': 'Owner',
            'period': 'September 2008 \u2013 Present (17 years 7 months)',
            'bullets': ['K\u00f8ge Municipality'],
        },
        {
            'company': 'Home development',
            'title': 'Developer',
            'period': 'April 1983 \u2013 Present (43 years)',
            'bullets': [
                'Developed Game for C-64, demos for C-64 and Amiga and a task / thread / resource manager for Amiga OS.',
                'Developed music editor software C-64, Amiga and PC Dos.',
                'Currently developing Android application for learning purpose and application with Rust and Svelte',
            ],
        },
    ],
    'repos': [
        'https://github.com/paletteguy/profile',
        'https://github.com/Laerdal/vex-kernel-checker',
        'https://github.com/Laerdal/linux-fslc',
        'https://github.com/Laerdal/meta-dependencytrack',
        'https://github.com/Laerdal-Medical/simserver-imager',
    ],
    'education_heading': 'Software development',
    'education': [
        {
            'institution': 'Niels Brock',
            'degree': "Bachelor's Degree, Computer Science",
            'period': 'September 1990 \u2013 June 1992',
        },
        {
            'institution': 'Niels Brock',
            'degree': "Bachelor's Degree, Business/Commerce, General",
            'period': 'August 1988 \u2013 June 1990',
        },
        {
            'institution': 'EFG Handel og Kontor',
            'degree': 'Basic business school',
            'period': 'August 1987 \u2013 June 1988',
        },
        {
            'institution': 'Krogaardskolen',
            'degree': 'High school diploma',
            'period': '1976 \u2013 1987',
        },
    ],
}


# ============================================================
# DOCX GENERATION
# ============================================================

def generate_docx(data, photo_path):
    doc = Document()

    # -- Style defaults --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)

    # -- Helper functions --
    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x62)
        return h

    def add_experience(company, title, period, description_lines):
        p = doc.add_paragraph()
        run = p.add_run(company)
        run.bold = True
        run.font.size = Pt(12)
        p.space_after = Pt(0)

        p2 = doc.add_paragraph()
        run2 = p2.add_run(title)
        run2.italic = True
        run2.font.size = Pt(10.5)
        p2.space_after = Pt(0)

        p3 = doc.add_paragraph()
        run3 = p3.add_run(period)
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p3.space_after = Pt(4)

        for line in description_lines:
            p4 = doc.add_paragraph(line, style='List Bullet')
            p4.paragraph_format.space_after = Pt(2)

        # Small spacer
        doc.add_paragraph().paragraph_format.space_before = Pt(4)

    # -- Header: name --
    h = doc.add_heading(data['name'], level=0)
    h.paragraph_format.space_after = Pt(0)
    h.paragraph_format.space_before = Pt(0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x62)
        run.font.size = Pt(28)
    # Remove the bottom border that Title style adds
    pPr = h._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is not None:
        pPr.remove(pBdr)
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        el = pBdr.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto',
        })
        pBdr.append(el)
    pPr.append(pBdr)

    # -- Tagline --
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(data['tagline'])
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)

    # -- Contact table (left: contact, right: photo) --
    photo_width = Inches(1.4)
    photo_width_twips = str(int(1.4 * 1440))
    left_width_twips = str(int((6.5 - 1.4) * 1440))
    page_width_twips = str(int(6.5 * 1440))

    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()

    # Remove any table style
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)

    # Fixed table layout
    tblLayout = tblPr.makeelement(qn('w:tblLayout'), {qn('w:type'): 'fixed'})
    tblPr.append(tblLayout)

    # Explicit table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = tblPr.makeelement(qn('w:tblW'), {})
        tblPr.append(tblW)
    tblW.set(qn('w:w'), page_width_twips)
    tblW.set(qn('w:type'), 'dxa')

    # Table borders: none
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto',
        })
        borders.append(el)
    tblPr.append(borders)

    # Grid columns
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = tbl.makeelement(qn('w:tblGrid'), {})
    gridCol1 = tblGrid.makeelement(qn('w:gridCol'), {qn('w:w'): left_width_twips})
    gridCol2 = tblGrid.makeelement(qn('w:gridCol'), {qn('w:w'): photo_width_twips})
    tblGrid.append(gridCol1)
    tblGrid.append(gridCol2)
    tblPr.addnext(tblGrid)

    # Cell widths and borders
    for i, cell in enumerate(header_table.rows[0].cells):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = tcPr.makeelement(qn('w:tcW'), {})
            tcPr.append(tcW)
        tcW.set(qn('w:w'), left_width_twips if i == 0 else photo_width_twips)
        tcW.set(qn('w:type'), 'dxa')
        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        for edge_name in ('top', 'left', 'bottom', 'right'):
            el = tcBorders.makeelement(qn(f'w:{edge_name}'), {
                qn('w:val'): 'none', qn('w:sz'): '0',
                qn('w:space'): '0', qn('w:color'): 'auto',
            })
            tcBorders.append(el)
        old = tcPr.find(qn('w:tcBorders'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcBorders)

    # Left cell: contact details
    left_cell = header_table.cell(0, 0)
    left_cell.paragraphs[0].clear()
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    contact_lines = [('Contact details', ''), ('', '')] + data['contact']
    for i, (label, value) in enumerate(contact_lines):
        if i > 0:
            p.add_run('\n')
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_value = p.add_run(value)
        run_value.font.size = Pt(10)

    # Right cell: photo
    right_cell = header_table.cell(0, 1)
    right_cell.paragraphs[0].clear()
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    right_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    img = Image.open(photo_path)
    img.thumbnail((350, 350))
    img_stream = io.BytesIO()
    img.save(img_stream, format='PNG')
    img_stream.seek(0)
    right_cell.paragraphs[0].add_run().add_picture(img_stream, width=photo_width)

    # -- Summary --
    add_heading_styled('Summary', level=1)
    doc.add_paragraph(data['summary'])

    # -- Core Skills --
    add_heading_styled('Core Skills', level=1)
    doc.add_paragraph(data['skills'])

    # -- Languages --
    add_heading_styled('Languages', level=1)
    for lang, desc in data['languages']:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(lang + ': ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # -- Experience --
    add_heading_styled('Experience', level=1)
    for exp in data['experience']:
        add_experience(exp['company'], exp['title'], exp['period'], exp['bullets'])
        if exp.get('page_break_after'):
            doc.add_page_break()

    doc.add_page_break()

    # -- Github Repositories --
    add_heading_styled('Github repositories', level=1)
    for repo in data['repos']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(repo)

    # -- Education --
    add_heading_styled('Education', level=1)

    p = doc.add_paragraph()
    run = p.add_run(data['education_heading'])
    run.bold = True

    for edu in data['education']:
        p = doc.add_paragraph()
        run = p.add_run(edu['institution'])
        run.bold = True
        p.add_run('\n' + edu['degree'])
        p.add_run('\n' + edu['period'])

    return doc


# ============================================================
# MARKDOWN GENERATION
# ============================================================

def generate_md(data):
    lines = []

    # Header table (photo left, contact right)
    lines.append('<table>')
    lines.append('<tr>')
    lines.append('<td width="120" valign="top">')
    lines.append('')
    lines.append(f'![{data["name"].split()[0]} {data["name"].split()[-1]}]({data["photo"]})')
    lines.append('')
    lines.append('</td>')
    lines.append('<td>')
    lines.append('')
    lines.append(f'# {data["name"]}')
    lines.append('')
    lines.append(f'**{data["tagline"]}**')
    lines.append('')
    lines.append('**Contact details**<br>')
    lines.append('<br>')
    for i, (label, value) in enumerate(data['contact']):
        suffix = '<br>' if i < len(data['contact']) - 1 else ''
        lines.append(f'**{label}** {value}{suffix}')
    lines.append('')
    lines.append('</td>')
    lines.append('</tr>')
    lines.append('</table>')
    lines.append('')
    lines.append('---')
    lines.append('')

    # Summary
    lines.append('## Summary')
    lines.append('')
    lines.append(data['summary'])
    lines.append('')
    lines.append('---')
    lines.append('')

    # Core Skills
    lines.append('## Core Skills')
    lines.append('')
    lines.append(data['skills'])
    lines.append('')
    lines.append('---')
    lines.append('')

    # Languages
    lines.append('## Languages')
    lines.append('')
    for lang, desc in data['languages']:
        lines.append(f'**{lang}:** {desc}')
        lines.append('')
    lines.append('---')
    lines.append('')

    # Experience
    lines.append('## Experience')
    lines.append('')
    for exp in data['experience']:
        lines.append(f'### {exp["company"]}')
        lines.append(f'*{exp["title"]}*')
        lines.append(exp['period'])
        lines.append('')
        for bullet in exp['bullets']:
            lines.append(f'- {bullet}')
        lines.append('')
    lines.append('---')
    lines.append('')

    # Github repositories
    lines.append('## Github repositories')
    lines.append('')
    for repo in data['repos']:
        lines.append(f'- {repo}')
    lines.append('')
    lines.append('---')
    lines.append('')

    # Education
    lines.append('## Education')
    lines.append('')
    lines.append(f'**{data["education_heading"]}**')
    lines.append('')
    for edu in data['education']:
        lines.append(f'**{edu["institution"]}**')
        lines.append(edu['degree'])
        lines.append(edu['period'])
        lines.append('')

    return '\n'.join(lines)


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    photo_path = os.path.join(script_dir, 'me.png')

    # Generate .docx
    docx_name = 'Karsten Opdal - CV.docx'
    doc = generate_docx(CV_DATA, photo_path)
    doc.save(docx_name)
    print(f'Saved to: {docx_name}')

    # Generate README.md
    md_name = os.path.join(script_dir, 'README.md')
    md_content = generate_md(CV_DATA)
    with open(md_name, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f'Saved to: {md_name}')

    # Convert to PDF via LibreOffice
    output_dir = os.path.dirname(os.path.abspath(docx_name)) or '.'
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', output_dir, docx_name
    ], check=True)
    pdf_name = os.path.splitext(docx_name)[0] + '.pdf'
    print(f'Saved to: {pdf_name}')
